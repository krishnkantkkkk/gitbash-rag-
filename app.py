import hashlib
import os
import pickle
import io
import fitz
import faiss
import docx
import tempfile
import base64
import ollama
from flask import Flask, Response, render_template, request, jsonify, stream_with_context, json
from sentence_transformers import SentenceTransformer, CrossEncoder
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_ollama.chat_models import ChatOllama
from langchain.text_splitter import RecursiveCharacterTextSplitter
from faster_whisper import WhisperModel
import pytesseract
from PIL import Image
import numpy as np
from concurrent.futures import ThreadPoolExecutor
import threading
from transformers import BlipProcessor, BlipForConditionalGeneration
from PIL import Image

# --- App Initialization ---
app = Flask(__name__, static_folder='temp', static_url_path='/temp')

# --- Global Variables & Model Loading ---
all_documents_metadata = []
vector_store = None
session_uploaded_files = set()
session_file_hashes = {}
session_file_indices = {}
CACHE_DIR = "cache"

# ✨ OPTIMIZATION: Lazy loading of models to reduce startup time
_embedding_model = None
_reranker = None
_whisper_model = None
_ex_llm = None
_llm = None
_model_lock = threading.Lock()

def get_embedding_model():
    global _embedding_model
    if _embedding_model is None:
        with _model_lock:
            if _embedding_model is None:
                _embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
    return _embedding_model

def get_reranker():
    global _reranker
    if _reranker is None:
        with _model_lock:
            if _reranker is None:
                _reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')
    return _reranker

def get_whisper_model():
    global _whisper_model
    if _whisper_model is None:
        with _model_lock:
            if _whisper_model is None:
                _whisper_model = WhisperModel("base.en", device="cpu", compute_type="int8")
    return _whisper_model

def get_llm():
    global _llm
    if _llm is None:
        with _model_lock:
            if _llm is None:
                _llm = ChatOllama(model="llama3.2")
    return _llm

def get_ex_llm():
    global _ex_llm
    if _ex_llm is None:
        with _model_lock:
            if _ex_llm is None:
                _ex_llm = ChatOllama(model="gemma3:1b")
    return _ex_llm

executor = ThreadPoolExecutor(max_workers=4)

# --- Processing Functions ---

def describe_image_with_vision_model(image_path):
    try:
        print("Reading Image...")
        processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-large")
        model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-large")
        image = Image.open(image_path).convert("RGB")
        inputs = processor(image, return_tensors="pt")

        outputs = model.generate(
            **inputs,
            max_length=100,
            num_return_sequences=5,
            do_sample=True,
            top_k=50,
            top_p=0.95
        )
        caption = "Here are some descriptions of the image: "
        for out in outputs:
            caption += processor.decode(out, skip_special_tokens=True) + " "
        print(caption)
        return caption
    except Exception as e:
        print(f"Vision model description failed: {e}")
        return "Unable to generate detailed image description."


def process_pdf(file_storage):
    """Processes a PDF file by extracting text chunks and images with OCR + Vision descriptions."""
    file_bytes = io.BytesIO(file_storage.read())
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    
    processed_data = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)

    def process_page(page_num):
        page = doc[page_num]
        page_data = []
        
        # Process text
        text = page.get_text()
        if text.strip():
            chunks = text_splitter.split_text(text)
            for chunk in chunks:
                page_data.append({
                    "text": chunk,
                    "page_num": page_num + 1,
                    "source_filename": file_storage.filename,
                    "type": "text"
                })

        # Process images with vision model
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            try:
                base_image = doc.extract_image(xref)
                img_bytes = base_image["image"]
                pil_image = Image.open(io.BytesIO(img_bytes)).convert("RGB")

                os.makedirs("temp", exist_ok=True)
                ext = base_image["ext"]
                img_filename = f"{os.path.splitext(file_storage.filename)[0]}_p{page_num+1}_{img_index}.{ext}"
                img_path = os.path.join("temp", img_filename)
                pil_image.save(img_path)

                # OCR for text extraction
                ocr_text = ""
                if pil_image.width > 50 and pil_image.height > 50:
                    ocr_text = pytesseract.image_to_string(pil_image)

                vision_description = describe_image_with_vision_model(img_path)
                
                # Create rich image document
                image_doc_text = f"""Image from page {page_num + 1} of {file_storage.filename}
Vision Description: {vision_description}
OCR Text: {ocr_text.strip()}""".strip()

                page_data.append({
                    "text": image_doc_text,
                    "image_path": img_filename,
                    "page_num": page_num + 1,
                    "source_filename": file_storage.filename,
                    "type": "image",
                })
            except Exception as e:
                print(f"Warning: Could not process image {img_index} on page {page_num+1}: {e}")
        
        return page_data

    # Process pages in parallel
    with ThreadPoolExecutor(max_workers=4) as page_executor:
        page_results = list(page_executor.map(process_page, range(len(doc))))
    
    for page_data in page_results:
        processed_data.extend(page_data)

    doc.close()
    return processed_data

def process_docx(file_storage):
    """Processes a DOCX file by extracting content in proper sequence (text and images)."""
    file_bytes = io.BytesIO(file_storage.read())
    doc = docx.Document(file_bytes)
    
    doc_data = []
    os.makedirs("temp", exist_ok=True)
    
    # Create a mapping of image relationships
    image_rels = {}
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_rels[rel.rId] = rel
    
    # Track position for sequential processing
    position = 0
    current_text = ""
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    
    # First pass: collect all elements with their types in sequence
    all_elements = []
    for element in doc.element.body:
        if element.tag.endswith('p'):
            # Check for inline images first
            has_image = False
            for run in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                for drawing in run.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
                    blip = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    if blip is not None:
                        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed and embed in image_rels:
                            # Add any text before the image
                            para_text = ""
                            for para in doc.paragraphs:
                                if para._element == element:
                                    para_text = para.text
                                    break
                            if para_text.strip():
                                all_elements.append({"type": "text", "content": para_text})
                            
                            # Add the image
                            all_elements.append({"type": "image", "embed": embed, "rel": image_rels[embed]})
                            has_image = True
            
            # If no image in this paragraph, just add text
            if not has_image:
                para_text = ""
                for para in doc.paragraphs:
                    if para._element == element:
                        para_text = para.text
                        break
                
                if para_text.strip():
                    all_elements.append({"type": "text", "content": para_text})
    
    # Second pass: process elements with context
    for idx, elem_info in enumerate(all_elements):
        if elem_info["type"] == "text":
            para_text = elem_info["content"]
            current_text += para_text + "\n"
            
            # Check if we should chunk the accumulated text
            if len(current_text) >= 1000:
                chunks = text_splitter.split_text(current_text.strip())
                for chunk in chunks:
                    position += 1
                    doc_data.append({
                        "text": chunk,
                        "page_num": position,
                        "source_filename": file_storage.filename,
                        "type": "text"
                    })
                current_text = ""
        
        elif elem_info["type"] == "image":
            # Flush any accumulated text before adding image
            context_before = ""
            if current_text.strip():
                chunks = text_splitter.split_text(current_text.strip())
                for chunk in chunks:
                    position += 1
                    doc_data.append({
                        "text": chunk,
                        "page_num": position,
                        "source_filename": file_storage.filename,
                        "type": "text"
                    })
                # Keep recent text as context for the image
                context_before = current_text[-500:] if len(current_text) > 500 else current_text
                current_text = ""
            else:
                # Use previous chunks as context
                for prev_doc in reversed(doc_data[-3:]):  # Last 3 chunks
                    if prev_doc['type'] == 'text':
                        context_before = prev_doc['text'][:300] + " " + context_before
                context_before = context_before.strip()[:500]
            
            # Look ahead for context after the image
            context_after = ""
            for next_idx in range(idx + 1, min(idx + 4, len(all_elements))):
                if all_elements[next_idx]["type"] == "text":
                    context_after += all_elements[next_idx]["content"] + " "
                    if len(context_after) >= 500:
                        break
            context_after = context_after.strip()[:500]
            
            # Combine before and after context
            full_context = f"{context_before}\n[IMAGE HERE]\n{context_after}".strip()
            
            # Process the image
            try:
                rel = elem_info["rel"]
                img_data = rel.target_part.blob
                img = Image.open(io.BytesIO(img_data)).convert("RGB")
                
                ext = rel.target_ref.split('.')[-1]
                img_filename = f"{os.path.splitext(file_storage.filename)[0]}_img{position}.{ext}"
                img_path = os.path.join("temp", img_filename)
                img.save(img_path)
                
                # OCR
                ocr_text = ""
                if img.width > 50 and img.height > 50:
                    ocr_text = pytesseract.image_to_string(img)
                
                # Vision description with surrounding context
                vision_description = describe_image_with_vision_model(img_path)
                
                image_doc_text = f"""Image from {file_storage.filename}
Vision Description: {vision_description}
OCR Text: {ocr_text.strip()}""".strip()
                
                position += 1
                doc_data.append({
                    "text": image_doc_text,
                    "image_path": img_filename,
                    "page_num": position,
                    "source_filename": file_storage.filename,
                    "type": "image",
                })
            except Exception as e:
                print(f"Warning: Could not process DOCX image at position {position}: {e}")
    
    # Add any remaining text
    if current_text.strip():
        chunks = text_splitter.split_text(current_text.strip())
        for chunk in chunks:
            position += 1
            doc_data.append({
                "text": chunk,
                "page_num": position,
                "source_filename": file_storage.filename,
                "type": "text"
            })
    
    return doc_data

def process_audio(file_storage):
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_storage.filename.split('.')[-1]}") as temp_audio:
        file_storage.save(temp_audio.name)
        temp_audio_path = temp_audio.name

    try:
        whisper = get_whisper_model()
        segments, _ = whisper.transcribe(temp_audio_path, beam_size=5)
        timed_segments = []
        for seg in segments:
            timed_segments.append({
                "text": seg.text.strip(),
                "start": seg.start,
                "end": seg.end
            })
    finally:
        os.remove(temp_audio_path)

    if not timed_segments:
        return []

    # Smart chunking with timing
    chunk_data = []
    current_chunk = {
        "text": "",
        "start_time": None,
        "end_time": None,
        "segments": []
    }
    
    current_length = 0
    max_chunk_size = 1000
    overlap_size = 150
    
    for segment in timed_segments:
        segment_text = segment["text"]
        segment_length = len(segment_text)
        
        if current_length + segment_length > max_chunk_size and current_chunk["text"]:
            chunk_data.append({
                "text": current_chunk["text"].strip(),
                "source_filename": file_storage.filename,
                "type": "audio",
                "start_time": current_chunk["start_time"],
                "end_time": current_chunk["end_time"],
                "duration": current_chunk["end_time"] - current_chunk["start_time"],
                "page_num": len(chunk_data) + 1
            })
            
            overlap_text = ""
            overlap_length = 0
            for prev_seg in reversed(current_chunk["segments"]):
                if overlap_length + len(prev_seg["text"]) <= overlap_size:
                    overlap_text = prev_seg["text"] + " " + overlap_text
                    overlap_length += len(prev_seg["text"])
                else:
                    break
            
            current_chunk = {
                "text": overlap_text + " " + segment_text if overlap_text else segment_text,
                "start_time": segment["start"],
                "end_time": segment["end"],
                "segments": [segment]
            }
            current_length = len(current_chunk["text"])
        else:
            if current_chunk["start_time"] is None:
                current_chunk["start_time"] = segment["start"]
            current_chunk["end_time"] = segment["end"]
            
            if current_chunk["text"]:
                current_chunk["text"] += " " + segment_text
            else:
                current_chunk["text"] = segment_text
                
            current_chunk["segments"].append(segment)
            current_length = len(current_chunk["text"])
    
    if current_chunk["text"]:
        chunk_data.append({
            "text": current_chunk["text"].strip(),
            "source_filename": file_storage.filename,
            "type": "audio",
            "start_time": current_chunk["start_time"],
            "end_time": current_chunk["end_time"],
            "duration": current_chunk["end_time"] - current_chunk["start_time"],
            "page_num": len(chunk_data) + 1
        })
        
    return chunk_data

def create_vector_store_from_docs(documents):
    embedding_model = get_embedding_model()
    texts = [doc['text'] for doc in documents]
    embeddings = embedding_model.encode(texts, convert_to_tensor=True, show_progress_bar=False)
    
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings.cpu().numpy())
    return index
    

def format_timestamp(seconds):
    if seconds is None:
        return "00:00"
    minutes, secs = divmod(int(seconds), 60)
    return f"{minutes:02d}:{secs:02d}"


def expand_query(query):
    print("Expanding query for better retrieval...")
    template = """Based on the user's question, generate 3 additional, different, and more specific queries that are likely to find relevant documents in a vector database.
Focus on rephrasing, using synonyms, and breaking down the question into sub-questions.
Provide ONLY the queries, each on a new line. Do not number them or add any other text. And just questions no explainations, nothing else.

Original Question: {question}

Generated Queries:"""
    
    llm = get_ex_llm()
    prompt = ChatPromptTemplate.from_template(template)
    chain = prompt | llm | StrOutputParser()
    
    try:
        response = chain.invoke({"question": query})
        expanded_queries = [q.strip() for q in response.strip().split('\n') if q.strip()]
        all_queries = [query] + expanded_queries[:3]
        print("Expanded Queries:", all_queries)
        return list(set(all_queries))
    except Exception as e:
        print(f"Query expansion failed: {e}")
        return [query]


def load_from_cache(file_hash):
    """Load processed document data and embeddings from cache."""
    cache_path = os.path.join(CACHE_DIR, file_hash)
    docs_path = os.path.join(cache_path, "documents.pkl")
    embeddings_path = os.path.join(cache_path, "embeddings.npy")
    
    if os.path.exists(docs_path):
        try:
            with open(docs_path, "rb") as f:
                docs = pickle.load(f)
            
            embeddings = None
            if os.path.exists(embeddings_path):
                embeddings = np.load(embeddings_path)
                
            return {
                "docs": docs,
                "embeddings": embeddings
            }
        except Exception as e:
            print(f"Could not load cache for {file_hash}: {e}")
    return None


def save_to_cache(file_hash, docs, embeddings=None):
    """Save processed document data and embeddings to cache."""
    cache_path = os.path.join(CACHE_DIR, file_hash)
    os.makedirs(cache_path, exist_ok=True)
    
    with open(os.path.join(cache_path, "documents.pkl"), "wb") as f:
        pickle.dump(docs, f)
    
    if embeddings is not None:
        np.save(os.path.join(cache_path, "embeddings.npy"), embeddings)


# --- Flask Routes ---

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    global all_documents_metadata, vector_store, session_uploaded_files, session_file_hashes, session_file_indices
    try:
        if 'files' not in request.files:
            return jsonify({'error': 'No files provided'}), 400

        files = request.files.getlist('files')
        newly_processed_docs = []
        processed_filenames = []
        
        all_new_embeddings = []
        all_new_docs = []

        for file in files:
            file_bytes = file.read()
            file.seek(0)
            file_hash = hashlib.sha256(file_bytes).hexdigest()
            
            if file.filename in session_uploaded_files and session_file_hashes.get(file.filename) == file_hash:
                print(f"File {file.filename} already uploaded, skipping...")
                continue
            
            cached_data = load_from_cache(file_hash)
            
            if cached_data is not None and cached_data["docs"] is not None:
                print(f"Loading {file.filename} from cache...")
                docs = cached_data["docs"]
                cached_embeddings = cached_data["embeddings"]
                
                for doc in docs:
                    doc['source_filename'] = file.filename
                
                if cached_embeddings is not None:
                    print(f"Using cached embeddings for {file.filename}")
                    new_embeddings = cached_embeddings
                else:
                    print(f"Creating embeddings for cached documents of {file.filename}")
                    embedding_model = get_embedding_model()
                    new_texts = [doc['text'] for doc in docs]
                    new_embeddings = embedding_model.encode(new_texts, convert_to_tensor=True, show_progress_bar=False).cpu().numpy()
                    save_to_cache(file_hash, docs, new_embeddings)
                    
            else:
                print(f"Processing new file: {file.filename}")
                
                filename = file.filename.lower()
                if filename.endswith('.pdf'):
                    docs = process_pdf(file)
                elif filename.endswith('.docx'):
                    docs = process_docx(file)
                elif filename.endswith(('.mp3', '.wav', '.m4a', '.ogg')):
                    docs = process_audio(file)
                else:
                    continue
                
                if not docs: 
                    continue
                
                new_texts = [doc['text'] for doc in docs]
                embedding_model = get_embedding_model()
                new_embeddings = embedding_model.encode(new_texts, convert_to_tensor=True, show_progress_bar=False).cpu().numpy()
                
                save_to_cache(file_hash, docs, new_embeddings)
            
            start_idx = len(all_documents_metadata) + len(all_new_docs)
            end_idx = start_idx + len(docs)
            session_file_indices[file.filename] = {
                "start": start_idx,
                "end": end_idx,
                "count": len(docs)
            }
            
            all_new_docs.extend(docs)
            all_new_embeddings.append(new_embeddings)
            session_uploaded_files.add(file.filename)
            session_file_hashes[file.filename] = file_hash
            processed_filenames.append(file.filename)

        if all_new_embeddings:
            combined_embeddings = np.vstack(all_new_embeddings)
            
            if vector_store is None:
                dimension = combined_embeddings.shape[1]
                vector_store = faiss.IndexFlatL2(dimension)
            
            vector_store.add(combined_embeddings)
            all_documents_metadata.extend(all_new_docs)
            
            print(f"Added {len(all_new_docs)} total chunks from {len(processed_filenames)} files")

        return jsonify({'message': 'Files processed successfully', 'filenames': processed_filenames})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/ask', methods=['POST'])
def ask_question():
    if vector_store is None or vector_store.ntotal == 0: 
        return jsonify({'error': 'No documents uploaded yet'}), 400
    
    data = request.get_json()
    question = data.get('question')
    if not question: 
        return jsonify({'error': 'No question provided'}), 400

    queries = [question]
    
    queries = expand_query(question)
    
    embedding_model = get_embedding_model()
    query_embeddings = embedding_model.encode(queries, convert_to_tensor=True, show_progress_bar=False).cpu().numpy()

    k_retrieval = 5
    distances, ids = vector_store.search(query_embeddings, k_retrieval)
    
    unique_ids = set()
    for id_list in ids:
        for i in id_list:
            if i != -1 and 0 <= i < len(all_documents_metadata):
                unique_ids.add(i)
    
    candidate_docs = [all_documents_metadata[i] for i in unique_ids]
    
    if not candidate_docs:
        return Response(stream_with_context(iter(["<div><p>I couldn't find any relevant information in the uploaded documents to answer your question.</p></div>"])))

    # Rerank candidates
    reranker = get_reranker()
    rerank_pairs = [[question, doc['text']] for doc in candidate_docs]
    scores = reranker.predict(rerank_pairs)
    
    doc_scores = list(zip(candidate_docs, scores))
    doc_scores.sort(key=lambda x: x[1], reverse=True)
    
    # ✨ Increase top_k for image queries
    top_k_reranked = 3 if any(keyword in question.lower() for keyword in ['image', 'graph', 'chart', 'diagram', 'picture', 'show']) else 2
    retrieved_results = doc_scores[:top_k_reranked]
    retrieved_docs_metadata = [res[0] for res in retrieved_results]
    
    context_text = "\n\n".join([f"Source from {doc['source_filename']}, Page/Chunk {doc.get('page_num', 'N/A')}:\n{doc['text']}" for doc in retrieved_docs_metadata])
    
    # ✨ Enhanced prompt for image handling
    template = """Answer the question based ONLY on the following context.
Context:
{context}

Your answer must be in HTML format, enclosed within a single <div> tag. Do not use markdown, backticks, or any styling. If the question asks to "show" an image, graph, chart, or diagram, 
then you have to just one line caption it if available but don't use <img> tag for it. if there is nothing to show, just answer normally. If the context does not contain the answer, say "I couldn't find any relevant information in the uploaded documents to answer your question." Do not make up answers. Be concise.

Question: {question}"""
    
    prompt = ChatPromptTemplate.from_template(template)
    llm = get_llm()
    rag_chain = prompt | llm | StrOutputParser()
    def generate():
        full_response = ""
        for chunk in rag_chain.stream({"context": context_text, "question": question}):
            full_response += chunk
            yield chunk
        # ✨ Check if we should display images
        should_show_images = any(keyword in question.lower() for keyword in ['show', 'display', 'image', 'graph', 'chart', 'diagram', 'picture'])
        
        sources = []
        for doc, score in retrieved_results:
            source_obj = {
                "source_filename": doc['source_filename'],
                "page_num": doc.get('page_num', 0),
                "source_content": doc['text'],
                "type": doc.get('type', 'unknown'),
                "score": float(score)
            }
            if doc.get('type') == 'image': 
                source_obj['image_path'] = doc['image_path']
                source_obj['vision_description'] = doc.get('vision_description', '')
                source_obj['show_inline'] = should_show_images
            
            if doc.get('type') == 'audio':
                source_obj['start_time'] = doc.get('start_time')
                source_obj['end_time'] = doc.get('end_time')
                source_obj['duration'] = doc.get('duration')
                source_obj['timestamp_display'] = f"{format_timestamp(doc.get('start_time'))} - {format_timestamp(doc.get('end_time'))}"
            
            sources.append(source_obj)
        yield json.dumps({"type": "sources", "content": sources})

    return Response(stream_with_context(generate()), mimetype='text/plain')

@app.route('/transcribe', methods=['POST'])
def transcribe_audio():
    """Endpoint to transcribe spoken audio from the frontend."""
    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400
        
        audio_file = request.files['audio']
        with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_audio:
            audio_file.save(temp_audio.name)
            temp_audio_path = temp_audio.name
        
        try:
            whisper = get_whisper_model()
            segments, _ = whisper.transcribe(temp_audio_path, beam_size=5)
            transcription = " ".join([segment.text for segment in segments])
        finally:
            os.remove(temp_audio_path)
            
        return jsonify({'transcription': transcription})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/files', methods=['GET'])
def list_files():
    """Returns detailed info about files uploaded in the current session."""
    file_info = []
    for filename in sorted(session_uploaded_files):
        info = {
            "filename": filename,
            "hash": session_file_hashes.get(filename),
            "indices": session_file_indices.get(filename, {}),
            "chunk_count": session_file_indices.get(filename, {}).get("count", 0)
        }
        file_info.append(info)
    
    return jsonify({
        'files': list(session_uploaded_files),
        'detailed_info': file_info,
        'total_chunks': len(all_documents_metadata),
        'vector_store_size': vector_store.ntotal if vector_store else 0
    })


@app.route('/session-info', methods=['GET'])
def session_info():
    """Get detailed information about the current session."""
    return jsonify({
        'uploaded_files': list(session_uploaded_files),
        'file_indices': session_file_indices,
        'total_documents': len(all_documents_metadata),
        'vector_store_size': vector_store.ntotal if vector_store else 0,
        'cache_stats': {
            'cached_files': len([d for d in os.listdir(CACHE_DIR) if os.path.isdir(os.path.join(CACHE_DIR, d))]) if os.path.exists(CACHE_DIR) else 0
        }
    })


@app.route('/clear-session', methods=['POST'])
def clear_session():
    """Clear all uploaded files from the current session."""
    global all_documents_metadata, vector_store, session_uploaded_files, session_file_hashes, session_file_indices
    
    all_documents_metadata = []
    vector_store = None
    session_uploaded_files.clear()
    session_file_hashes.clear()
    session_file_indices.clear()
    
    return jsonify({'message': 'Session cleared successfully'})


if __name__ == '__main__':
    app.run(debug=True, port=5000, threaded=True)
